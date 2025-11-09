"""
Document & Media Embedding Pipeline
Handles ingestion, parsing, chunking, and embedding of documents.
"""

import os
from typing import List, Dict, Optional, Any, Callable
from pathlib import Path
import json
from datetime import datetime
import logging
import traceback

# Import logging utility
try:
    from scripts.logging_config import setup_logger, write_to_log, get_current_log_file
except ImportError:
    # Fallback if logging_config not available
    # Use session state to get the same log file as streamlit_app.py
    try:
        import streamlit as st
        USE_STREAMLIT = True
    except ImportError:
        USE_STREAMLIT = False
    
    def get_current_log_file():
        """Get current log file from session state or create one."""
        if USE_STREAMLIT and 'log_file_path' in st.session_state:
            return st.session_state.log_file_path
        else:
            # Fallback: create timestamped file
            os.makedirs("logs", exist_ok=True)
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            log_file = f"logs/logger_{timestamp}.log"
            if USE_STREAMLIT and 'log_file_path' not in st.session_state:
                st.session_state.log_file_path = log_file
            return log_file
    
    def setup_logger(name=None, level=logging.INFO):
        logger = logging.getLogger(name)
        if not logger.handlers:
            log_file = get_current_log_file()
            logging.basicConfig(
                level=level,
                format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                handlers=[
                    logging.FileHandler(log_file, encoding='utf-8', mode='a'),
                    logging.StreamHandler()
                ]
            )
        return logger
    
    def write_to_log(message: str, log_type: str = "INFO"):
        try:
            log_file = get_current_log_file()
            with open(log_file, "a", encoding="utf-8") as f:
                f.write(f"\n{'='*80}\n")
                f.write(f"{log_type} at {datetime.now().isoformat()}\n")
                f.write(f"{message}\n")
                f.write(f"{'='*80}\n\n")
        except:
            pass

logger = setup_logger(__name__)

# Document processing libraries
import PyPDF2
import docx

# LangChain imports - handle different versions
RecursiveCharacterTextSplitter = None
TextLoader = None
PyPDFLoader = None
Document = None

# Try importing text_splitter from different locations
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
except ImportError:
    try:
        from langchain_core.text_splitter import RecursiveCharacterTextSplitter
    except ImportError:
        try:
            # Fallback: create a simple text splitter
            class RecursiveCharacterTextSplitter:
                def __init__(self, chunk_size=512, chunk_overlap=50, length_function=len, separators=None):
                    self.chunk_size = chunk_size
                    self.chunk_overlap = chunk_overlap
                    self.length_function = length_function
                    self.separators = separators or ["\n\n", "\n", ". ", " ", ""]
                
                def split_documents(self, documents):
                    chunks = []
                    for doc in documents:
                        text = doc.page_content if hasattr(doc, 'page_content') else str(doc)
                        chunks.extend(self._split_text(text, doc.metadata if hasattr(doc, 'metadata') else {}))
                    return chunks
                
                def _split_text(self, text, metadata):
                    chunks = []
                    current_chunk = ""
                    for char in text:
                        current_chunk += char
                        if len(current_chunk) >= self.chunk_size:
                            chunks.append(Document(page_content=current_chunk, metadata=metadata))
                            current_chunk = current_chunk[-self.chunk_overlap:]
                    if current_chunk:
                        chunks.append(Document(page_content=current_chunk, metadata=metadata))
                    return chunks
        except:
            pass

# Try importing document loaders
try:
    from langchain_community.document_loaders import TextLoader, PyPDFLoader
except ImportError:
    try:
        from langchain.document_loaders import TextLoader, PyPDFLoader
    except ImportError:
        # Create simple loaders
        class TextLoader:
            def __init__(self, file_path, encoding='utf-8'):
                self.file_path = file_path
                self.encoding = encoding
            def load(self):
                with open(self.file_path, 'r', encoding=self.encoding) as f:
                    content = f.read()
                return [Document(page_content=content, metadata={"source": self.file_path})]
        
        class PyPDFLoader:
            def __init__(self, file_path):
                self.file_path = file_path
            def load(self):
                documents = []
                with open(self.file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        documents.append(Document(
                            page_content=text,
                            metadata={"page": page_num + 1, "source": self.file_path}
                        ))
                return documents

# Try importing Document class
try:
    from langchain_core.documents import Document
except ImportError:
    try:
        from langchain.schema import Document
    except ImportError:
        # Create simple Document class
        class Document:
            def __init__(self, page_content, metadata=None):
                self.page_content = page_content
                self.metadata = metadata or {}

# Try importing Ollama embeddings (primary) or OpenAI (fallback)
Embeddings = None
OllamaEmbeddings = None
OpenAIEmbeddings = None

try:
    from langchain_ollama import OllamaEmbeddings
    Embeddings = OllamaEmbeddings
    DEFAULT_EMBEDDING_MODEL = "llama3.1:latest"  # Use latest tag to match installed model
    EMBEDDING_DIMENSION = 4096  # Llama 3.1 embedding dimension
except ImportError:
    OllamaEmbeddings = None
    try:
        from langchain_openai import OpenAIEmbeddings
        Embeddings = OpenAIEmbeddings
        OpenAIEmbeddings = OpenAIEmbeddings  # Keep reference
        DEFAULT_EMBEDDING_MODEL = "text-embedding-ada-002"
        EMBEDDING_DIMENSION = 1536  # OpenAI ada-002 dimension
    except ImportError:
        try:
            from langchain_community.embeddings import OpenAIEmbeddings
            Embeddings = OpenAIEmbeddings
            OpenAIEmbeddings = OpenAIEmbeddings  # Keep reference
            DEFAULT_EMBEDDING_MODEL = "text-embedding-ada-002"
            EMBEDDING_DIMENSION = 1536
        except ImportError:
            Embeddings = None
            DEFAULT_EMBEDDING_MODEL = None
            EMBEDDING_DIMENSION = 1536

# Vector database
try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import (
        Distance, VectorParams, PointStruct, CollectionStatus,
        HnswConfigDiff, QuantizationConfig, ScalarQuantization, ScalarType
    )
except ImportError:
    # Qdrant not available, will use ChromaDB fallback
    QdrantClient = None
    Distance = None
    VectorParams = None
    PointStruct = None
    CollectionStatus = None
    HnswConfigDiff = None
    QuantizationConfig = None
    ScalarQuantization = None
    ScalarType = None

import streamlit as st


class DocumentIngestionPipeline:
    """
    Handles document ingestion, chunking, embedding, and storage.
    
    Design Decisions:
    - Uses Qdrant for vector storage (self-hosted, high performance, excellent metadata filtering)
    - Uses OpenAI embeddings (high quality, consistent with production expectations)
    - Implements context-aware chunking strategies per document type
    - Stores rich metadata for filtering and traceability
    """
    
    def __init__(self, project_id: str = "default", use_ollama: bool = True, vector_db_type: str = "Qdrant", 
                 embedding_model: str = "llama3.1", multimodal_embeddings=None):
        self.project_id = project_id
        self.use_ollama = use_ollama
        self.vector_db_type = vector_db_type  # "Qdrant" (default), "ChromaDB", or "Pinecone"
        self.embedding_model = embedding_model
        self.multimodal_embeddings = multimodal_embeddings
        
        # CRITICAL: Initialize vector database clients FIRST to ensure they always exist
        self.qdrant_client = None
        self.pinecone_index = None
        self.chroma_client = None
        self.chroma_collection = None
        
        # Initialize vector database based on user choice
        if self.vector_db_type == "Qdrant":
            self._init_qdrant()
        elif self.vector_db_type == "Pinecone":
            self._init_pinecone()
        else:  # Fallback to ChromaDB if Qdrant/Pinecone unavailable
            self._init_chroma_fallback()
        
        # Initialize embeddings - use multimodal if provided, otherwise text-only
        if self.multimodal_embeddings is not None:
            # Use multimodal embeddings
            self.embeddings = self.multimodal_embeddings
            self.embedding_dimension = self.multimodal_embeddings.get_dimension()
            st.info(f"✅ Using multimodal model: {embedding_model} ({self.embedding_dimension}D)")
        elif embedding_model not in ["llama3.1", "openai-ada-002"]:
            # Try to load multimodal model
            try:
                from scripts.multimodal_embeddings import MultimodalEmbeddings
                self.multimodal_embeddings = MultimodalEmbeddings(
                    model_name=embedding_model,
                    use_ollama=use_ollama
                )
                self.embeddings = self.multimodal_embeddings
                self.embedding_dimension = self.multimodal_embeddings.get_dimension()
                st.info(f"✅ Using multimodal model: {embedding_model} ({self.embedding_dimension}D)")
            except Exception as e:
                st.warning(f"Failed to load multimodal model {embedding_model}: {e}, falling back to text-only")
                self._init_text_embeddings()
        else:
            # Text-only embeddings
            self._init_text_embeddings()
    
    def _init_text_embeddings(self):
        """Initialize text-only embeddings - Ollama only, no OpenAI fallback."""
        if self.use_ollama and OllamaEmbeddings is not None:
            try:
                base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
                logger.info(f"Attempting to initialize Ollama embeddings with model {DEFAULT_EMBEDDING_MODEL} at {base_url}")
                self.embeddings = OllamaEmbeddings(
                    model=DEFAULT_EMBEDDING_MODEL,
                    base_url=base_url
                )
                # Test the connection by trying to embed a small test string
                # Skip the test if we're in a context where it might fail (e.g., during import)
                try:
                    _ = self.embeddings.embed_query("test")
                    logger.info(f"Successfully initialized Ollama with {DEFAULT_EMBEDDING_MODEL}")
                    self.embedding_dimension = EMBEDDING_DIMENSION
                    # Don't show success message here - it's shown by RAGRetrievalEngine
                except Exception as test_error:
                    # If test fails, log but don't fail completely - might be a temporary issue
                    logger.warning(f"Ollama connection test failed: {test_error}. Will try again on first use.")
                    # Still set dimension - the actual embedding will fail if Ollama is really down
                    self.embedding_dimension = EMBEDDING_DIMENSION
            except Exception as e:
                error_msg = str(e)
                logger.error(f"Ollama initialization failed: {error_msg}")
                if "Connection" in error_msg or "refused" in error_msg.lower() or "not responding" in error_msg.lower():
                    error_message = (
                        f"⚠️ Ollama server not accessible at {os.getenv('OLLAMA_BASE_URL', 'http://localhost:11434')}. "
                        f"Please ensure Ollama is running. You can start it by running 'ollama serve' in your terminal."
                    )
                    st.error(error_message)
                    raise Exception(error_message)
                else:
                    error_message = f"⚠️ Failed to initialize Ollama embeddings: {error_msg}. Please ensure Ollama is installed and running."
                    st.error(error_message)
                    raise Exception(error_message)
        else:
            if OllamaEmbeddings is None:
                raise Exception(
                    "Ollama embeddings not available. Please install langchain-ollama: pip install langchain-ollama"
                )
            else:
                raise Exception(
                    "Ollama is required but not enabled. Please set use_ollama=True or ensure Ollama is properly configured."
                )
    
    def _init_openai_embeddings(self):
        """Initialize OpenAI embeddings as fallback - DEPRECATED, not used anymore."""
        # This method is kept for backward compatibility but should not be called
        # We only use Ollama now
        raise Exception(
            "OpenAI embeddings are not supported. Please use Ollama instead. "
            "Ensure Ollama is installed and running: pip install langchain-ollama && ollama serve"
        )
    
    def _init_qdrant(self):
        """Initialize Qdrant client with connection health check."""
        logger.info("Initializing Qdrant client...")
        if QdrantClient is not None:
            qdrant_url = os.getenv("QDRANT_URL", "http://localhost:6333")
            qdrant_api_key = os.getenv("QDRANT_API_KEY", None)
            
            try:
                # Import health check utility (use importlib for modules starting with numbers)
                import importlib
                qdrant_health_module = importlib.import_module('scripts.08_utilities.qdrant_health')
                check_qdrant_connection = qdrant_health_module.check_qdrant_connection
                verify_qdrant_before_operation = qdrant_health_module.verify_qdrant_before_operation
                
                logger.debug(f"Checking Qdrant connection at {qdrant_url}")
                is_connected, error_msg, info = check_qdrant_connection(
                    qdrant_url=qdrant_url,
                    qdrant_api_key=qdrant_api_key,
                    timeout=5.0
                )
                
                if not is_connected:
                    logger.error(f"Qdrant connection check failed: {error_msg}")
                    st.warning(f"⚠️ {error_msg}")
                    self.qdrant_client = None
                    self._init_chroma_fallback()
                    return
                
                # Connection check passed, create client
                logger.debug(f"Connecting to Qdrant at {qdrant_url}")
                self.qdrant_client = QdrantClient(
                    url=qdrant_url,
                    api_key=qdrant_api_key,
                    timeout=30
                )
                
                # Verify connection with actual operation
                verify_qdrant_before_operation(self.qdrant_client, "Qdrant initialization")
                
                # Test connection and ensure collection
                self.qdrant_client.get_collections()
                self._ensure_collection()
                
                logger.info("Successfully connected to Qdrant")
                if info:
                    logger.info(f"Qdrant info: {info.get('collections_count', 0)} collections, "
                              f"response time: {info.get('response_time_ms', 'unknown')}ms")
                # Don't show success message here - it's shown by RAGRetrievalEngine
            except Exception as e:
                error_msg = f"Qdrant connection failed: {e}"
                logger.warning(error_msg)
                logger.debug(f"Qdrant error traceback: {traceback.format_exc()}")
                st.warning(f"{error_msg}. Falling back to ChromaDB.")
                st.info("💡 **To start Qdrant:** Run `docker run -p 6333:6333 qdrant/qdrant` or use `start_qdrant.bat`")
                self.qdrant_client = None
                self._init_chroma_fallback()
        else:
            logger.warning("Qdrant client not available (not installed)")
            st.warning("Qdrant not installed. Falling back to ChromaDB.")
            self._init_chroma_fallback()
    
    def _init_pinecone(self):
        """Initialize Pinecone client."""
        try:
            import pinecone
            from pinecone import Pinecone, ServerlessSpec
            
            api_key = os.getenv("PINECONE_API_KEY", "")
            if not api_key:
                st.error("Pinecone API key not found. Please set PINECONE_API_KEY or use the sidebar.")
                self._init_chroma_fallback()
                return
            
            # Initialize Pinecone
            pc = Pinecone(api_key=api_key)
            
            # Create or get index
            index_name = f"deep-intellect-rag-{self.project_id}".lower().replace("_", "-")
            
            # Check if index exists
            existing_indexes = [idx.name for idx in pc.list_indexes()]
            
            if index_name not in existing_indexes:
                # Get embedding dimension
                embedding_dim = getattr(self, 'embedding_dimension', EMBEDDING_DIMENSION)
                # Create new index
                pc.create_index(
                    name=index_name,
                    dimension=embedding_dim,
                    metric="cosine",
                    spec=ServerlessSpec(
                        cloud="aws",
                        region=os.getenv("PINECONE_ENVIRONMENT", "us-east-1")
                    )
                )
                st.success(f"✅ Created Pinecone index: {index_name} ({embedding_dim}D)")
            
            self.pinecone_index = pc.Index(index_name)
            st.success(f"✅ Connected to Pinecone index: {index_name}")
            
        except ImportError:
            st.error("Pinecone not installed. Install with: pip install pinecone-client")
            self._init_chroma_fallback()
        except Exception as e:
            st.error(f"Pinecone initialization failed: {e}. Falling back to ChromaDB.")
            self._init_chroma_fallback()
    
    def _ensure_collection(self):
        """Create Qdrant collection if it doesn't exist with performance optimizations."""
        if self.qdrant_client is None or VectorParams is None:
            return
            
        collection_name = f"project_{self.project_id}"
        
        try:
            collections = self.qdrant_client.get_collections().collections
            collection_names = [c.name for c in collections]
            
            if collection_name not in collection_names:
                # Get embedding dimension
                embedding_dim = getattr(self, 'embedding_dimension', EMBEDDING_DIMENSION)
                
                # ✅ Optimization 1: HNSW Index Configuration
                # m=32: Good balance between accuracy and speed (16-64 range)
                # ef_construct=200: Quality during index build (100-500 range)
                hnsw_config = None
                if HnswConfigDiff is not None:
                    hnsw_config = HnswConfigDiff(
                        m=32,  # Graph connectivity (higher = more accurate, slower build)
                        ef_construct=200  # Search quality during index build
                    )
                
                # ✅ Optimization 2: Scalar Quantization (int8)
                # Reduces memory footprint by ~4x and speeds up search
                # Use for datasets > 1M vectors or when RAM is limited
                quantization_config = None
                if QuantizationConfig is not None and ScalarQuantization is not None and ScalarType is not None:
                    # Check if we should enable quantization (for large collections or limited RAM)
                    # For now, disable quantization to avoid API compatibility issues
                    # quantization_config = QuantizationConfig(
                    #     scalar=ScalarQuantization(
                    #         type=ScalarType.INT8,
                    #         always_ram=True  # Keep quantized vectors in RAM for speed
                    #     )
                    # )
                    quantization_config = None  # Disabled due to API changes
                
                # Create collection with optimizations
                self.qdrant_client.create_collection(
                    collection_name=collection_name,
                    vectors_config=VectorParams(
                        size=embedding_dim,  # Embedding dimension (varies by model)
                        distance=Distance.COSINE,
                        hnsw_config=hnsw_config,  # HNSW optimization
                        quantization_config=quantization_config  # Quantization optimization
                    )
                )
                
                # ✅ Optimization 3: Index Metadata Fields for Fast Filtered Search
                # Index common filter fields to avoid full dataset scans
                self._index_metadata_fields(collection_name)
                
                st.success(f"✅ Created optimized collection: {collection_name} (HNSW m=32, ef_construct=200)")
                logger.info(f"Created Qdrant collection '{collection_name}' with performance optimizations")
            else:
                # Collection exists - check if we need to update it with optimizations
                self._update_collection_optimizations(collection_name)
                
        except Exception as e:
            logger.error(f"Failed to create collection: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            st.error(f"Failed to create collection: {e}")
    
    def _index_metadata_fields(self, collection_name: str):
        """Create indexes on metadata fields for fast filtered searches."""
        if self.qdrant_client is None:
            return
        
        try:
            # Common metadata fields used for filtering
            metadata_fields = [
                "file_name",
                "page_number",
                "section_number",
                "subsection_number",
                "cohort_id",
                "source_type",
                "speaker_role",
                "chunk_index"
            ]
            
            # Note: Qdrant automatically indexes payload fields, but we can optimize
            # by ensuring these fields are properly typed. For explicit indexing in newer
            # Qdrant versions, we would use create_payload_index, but for now we rely
            # on Qdrant's automatic indexing.
            logger.debug(f"Metadata fields for filtering: {metadata_fields}")
            
        except Exception as e:
            logger.warning(f"Failed to index metadata fields: {e}")
            # Non-critical, continue without explicit indexing
    
    def _update_collection_optimizations(self, collection_name: str):
        """Update existing collection with performance optimizations if not already applied."""
        if self.qdrant_client is None:
            return
        
        try:
            collection_info = self.qdrant_client.get_collection(collection_name)
            
            # Check if quantization is already enabled
            has_quantization = collection_info.config.params.quantization_config is not None
            
            # Check if HNSW is configured (it's enabled by default, but check params)
            has_hnsw = collection_info.config.params.vectors.hnsw_config is not None
            
            # Update quantization if not present (for existing collections)
            # Disabled due to API compatibility issues with newer Qdrant client versions
            # if not has_quantization and QuantizationConfig is not None and ScalarQuantization is not None and ScalarType is not None:
            #     try:
            #         quantization_config = QuantizationConfig(
            #             scalar=ScalarQuantization(
            #                 type=ScalarType.INT8,
            #                 always_ram=True
            #             )
            #         )
            #         self.qdrant_client.update_collection(
            #             collection_name=collection_name,
            #             quantization_config=quantization_config
            #         )
            #         logger.info(f"✅ Added quantization to existing collection: {collection_name}")
            #     except Exception as e:
            #         logger.warning(f"Could not add quantization to existing collection: {e}")
            pass  # Quantization disabled due to API changes
            
            # Update HNSW config if needed
            if not has_hnsw and HnswConfigDiff is not None:
                try:
                    hnsw_config = HnswConfigDiff(m=32, ef_construct=200)
                    # Note: HNSW config updates may require collection recreation
                    # For now, we log that optimization would be beneficial
                    logger.debug(f"HNSW optimization recommended for: {collection_name}")
                except Exception as e:
                    logger.debug(f"HNSW config update not available: {e}")
                    
        except Exception as e:
            logger.debug(f"Collection optimization check: {e}")
            # Non-critical, continue
    
    def _init_chroma_fallback(self):
        """Fallback to ChromaDB if Qdrant unavailable."""
        try:
            import chromadb
            from chromadb.config import Settings
            
            self.chroma_client = chromadb.Client(Settings(
                chroma_db_impl="duckdb+parquet",
                persist_directory="./chroma_db"
            ))
            self.chroma_collection = self.chroma_client.get_or_create_collection(
                name=f"project_{self.project_id}"
            )
            st.info("Using ChromaDB as fallback vector store.")
        except Exception as e:
            st.error(f"ChromaDB fallback also failed: {e}")
            self.chroma_client = None
    
    def is_file_processed(self, file_name: str) -> bool:
        """Check if a file has already been processed and stored."""
        file_name = Path(file_name).name
        
        try:
            if self.vector_db_type == "Qdrant" and self.qdrant_client:
                collection_name = f"project_{self.project_id}"
                # Check if collection exists
                collections = self.qdrant_client.get_collections().collections
                collection_names = [c.name for c in collections]
                
                if collection_name not in collection_names:
                    return False
                
                # Search for any chunks with this file name
                from qdrant_client.models import Filter, FieldCondition, MatchValue
                if Filter and FieldCondition and MatchValue:
                    file_filter = Filter(
                        must=[
                            FieldCondition(
                                key="file_name",
                                match=MatchValue(value=file_name)
                            )
                        ]
                    )
                    # Try to get at least one point with this file name
                    results = self.qdrant_client.scroll(
                        collection_name=collection_name,
                        scroll_filter=file_filter,
                        limit=1
                    )
                    return len(results[0]) > 0
                    
            elif self.vector_db_type == "Pinecone" and self.pinecone_index:
                # Check Pinecone for file
                # Note: Pinecone query requires vector, so we use metadata filter in search
                # For now, return False (can be enhanced later)
                return False
                
            elif hasattr(self, 'chroma_collection') and self.chroma_collection:
                # Check ChromaDB
                try:
                    results = self.chroma_collection.get(
                        where={"file_name": file_name},
                        limit=1
                    )
                    return len(results.get('ids', [])) > 0
                except:
                    return False
        except Exception as e:
            logger.warning(f"Error checking if file is processed: {e}")
            return False
        
        return False
    
    def ingest_file(
        self,
        file_path: str,
        file_type: str,
        source_type: str = "document",
        speaker_role: Optional[str] = None,
        cohort_id: Optional[str] = None,
        metadata: Optional[Dict] = None,
        progress_callback: Optional[Callable[[str, float, str], None]] = None,
        extract_images: bool = False
    ) -> Dict[str, Any]:
        """
        Ingest a file: parse, chunk, embed, and store.
        
        Args:
            file_path: Path to the file
            file_type: One of 'txt', 'pdf', 'docx', 'vtt', 'srt', 'json'
            source_type: Type of source (transcript, document, note, etc.)
            speaker_role: For transcripts (interviewer, interviewee, etc.)
            cohort_id: Participant grouping identifier
            metadata: Additional metadata
            
        Returns:
            Dict with ingestion results
        """
        logger.info(f"Starting ingestion: file_path={file_path}, file_type={file_type}, source_type={source_type}")
        logger.debug(f"Metadata: speaker_role={speaker_role}, cohort_id={cohort_id}, metadata={metadata}")
        
        # Progress callback helper
        def update_progress(step: str, progress: float, detail: str = ""):
            if progress_callback:
                progress_callback(step, progress, detail)
        
        try:
            update_progress("Checking file", 0.05, "Validating file...")
            
            # Check if file exists
            if not os.path.exists(file_path):
                error_msg = f"File not found: {file_path}"
                logger.error(error_msg)
                return {
                    "success": False,
                    "error": error_msg,
                    "message": error_msg
                }
            
            file_name = Path(file_path).name
            
            # Check if file already processed
            update_progress("Checking duplicates", 0.08, "Checking if file already exists...")
            if self.is_file_processed(file_name):
                logger.info(f"File {file_name} already processed, skipping")
                return {
                    "success": True,
                    "file_name": file_name,
                    "chunks_created": 0,
                    "chunks_stored": 0,
                    "message": f"File {file_name} already processed (skipped duplicate)",
                    "skipped": True
                }
            
            file_size = os.path.getsize(file_path)
            logger.info(f"File size: {file_size} bytes")
            
            update_progress("Parsing document", 0.15, f"Reading {file_type} file...")
            
            # Handle image files if multimodal mode
            if file_type in ["jpg", "jpeg", "png", "gif", "bmp", "webp"] and extract_images:
                documents = self._parse_image_file(file_path, file_type)
            else:
                # Parse document based on type
                logger.info(f"Parsing {file_type} file...")
                if file_type == "txt":
                    documents = self._parse_txt(file_path)
                elif file_type == "pdf":
                    documents = self._parse_pdf(file_path, extract_images=extract_images)
                elif file_type == "docx":
                    documents = self._parse_docx(file_path, extract_images=extract_images)
                elif file_type in ["vtt", "srt"]:
                    documents = self._parse_subtitle(file_path, file_type)
                elif file_type == "json":
                    documents = self._parse_json_transcript(file_path)
                else:
                    error_msg = f"Unsupported file type: {file_type}"
                    logger.error(error_msg)
                    raise ValueError(error_msg)
            
            logger.info(f"Parsed {len(documents)} document(s)")
            update_progress("Chunking documents", 0.30, f"Creating chunks from {len(documents)} document(s)...")
            
            # Chunk documents
            logger.info("Chunking documents...")
            chunks = self._chunk_documents(documents, file_type, source_type)
            logger.info(f"Created {len(chunks)} chunks")
            
            update_progress("Enriching metadata", 0.50, f"Adding metadata to {len(chunks)} chunks...")
            
            # Add metadata to chunks (file_name already set above)
            enriched_chunks = []
            logger.debug("Enriching chunks with metadata...")
            for i, chunk in enumerate(chunks):
                chunk_metadata = {
                    "file_name": file_name,
                    "source_type": source_type,
                    "project_id": self.project_id,
                    "chunk_index": i,
                    "created_at": datetime.now().isoformat(),
                    "word_count": len(chunk.page_content.split())
                }
                
                # Add image metadata if this chunk contains an image
                if hasattr(chunk, 'metadata') and chunk.metadata.get('has_image'):
                    chunk_metadata["has_image"] = True
                    chunk_metadata["image_path"] = chunk.metadata.get("image_path")
                    chunk_metadata["content_type"] = "image"
                    chunk_metadata["embedding_type"] = "multimodal"
                    if chunk.metadata.get("image_index") is not None:
                        chunk_metadata["image_index"] = chunk.metadata.get("image_index")
                    if chunk.metadata.get("page") is not None:
                        chunk_metadata["page"] = chunk.metadata.get("page")
                else:
                    chunk_metadata["embedding_type"] = "text"
                
                # Preserve page, section, and subsection metadata
                if hasattr(chunk, 'metadata'):
                    # Page number
                    if chunk.metadata.get("page") is not None:
                        chunk_metadata["page_number"] = chunk.metadata.get("page")
                    elif chunk.metadata.get("page_number") is not None:
                        chunk_metadata["page_number"] = chunk.metadata.get("page_number")
                    
                    # Section number
                    if chunk.metadata.get("section_number") is not None:
                        chunk_metadata["section_number"] = chunk.metadata.get("section_number")
                    
                    # Subsection number
                    if chunk.metadata.get("subsection_number") is not None:
                        chunk_metadata["subsection_number"] = chunk.metadata.get("subsection_number")
                
                if speaker_role:
                    chunk_metadata["speaker_role"] = speaker_role
                if cohort_id:
                    chunk_metadata["cohort_id"] = cohort_id
                if metadata:
                    chunk_metadata.update(metadata)
                
                # Add timestamp metadata for time-based media
                if hasattr(chunk, 'metadata') and 'start_time' in chunk.metadata:
                    chunk_metadata["start_time"] = chunk.metadata.get("start_time")
                    chunk_metadata["end_time"] = chunk.metadata.get("end_time")
                
                chunk.metadata = chunk_metadata
                enriched_chunks.append(chunk)
                
                # Update progress for metadata enrichment
                if (i + 1) % 10 == 0 or i == len(chunks) - 1:
                    progress = 0.50 + (0.20 * (i + 1) / len(chunks))
                    update_progress("Enriching metadata", progress, f"Processed {i + 1}/{len(chunks)} chunks")
            
            update_progress("Generating embeddings", 0.70, f"Creating embeddings for {len(enriched_chunks)} chunks...")
            
            # Generate embeddings and store
            logger.info(f"Storing {len(enriched_chunks)} chunks in {self.vector_db_type}...")
            stored_count = self._store_chunks(enriched_chunks, progress_callback=progress_callback)
            logger.info(f"Successfully stored {stored_count} chunks")
            
            update_progress("Complete", 1.0, f"Successfully stored {stored_count} chunks!")
            
            return {
                "success": True,
                "file_name": file_name,
                "chunks_created": len(enriched_chunks),
                "chunks_stored": stored_count,
                "message": f"Successfully ingested {file_name}"
            }
            
        except Exception as e:
            error_traceback = traceback.format_exc()
            error_msg = str(e)
            logger.error(f"Exception in ingest_file: {error_msg}")
            logger.error(f"Traceback:\n{error_traceback}")
            
            # Write detailed error to log file
            error_log_msg = f"""INGESTION ERROR
File: {file_path}
File Type: {file_type}
Source Type: {source_type}
Error: {error_msg}
Traceback:
{error_traceback}"""
            write_to_log(error_log_msg, "INGESTION ERROR")
            
            return {
                "success": False,
                "error": error_msg,
                "message": f"Failed to ingest file: {error_msg}",
                "traceback": error_traceback
            }
    
    def _parse_txt(self, file_path: str) -> List[Document]:
        """Parse plain text file."""
        try:
            loader = TextLoader(file_path, encoding='utf-8')
            return loader.load()
        except UnicodeDecodeError:
            # Fallback encoding
            loader = TextLoader(file_path, encoding='latin-1')
            return loader.load()
    
    def _parse_pdf(self, file_path: str, extract_images: bool = False) -> List[Document]:
        """Parse PDF file. Optionally extract images if extract_images=True."""
        documents = []
        
        try:
            loader = PyPDFLoader(file_path)
            loaded_docs = loader.load()
            # Ensure page numbers are preserved
            for doc in loaded_docs:
                if "page" not in doc.metadata:
                    # Try to extract from source or set default
                    doc.metadata["page"] = 1
                documents.extend(self._extract_sections_from_text(doc))
        except Exception as e:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    doc = Document(
                        page_content=text,
                        metadata={"page": page_num + 1, "page_number": page_num + 1}
                    )
                    # Extract sections from page text
                    documents.extend(self._extract_sections_from_text(doc))
        
        # Extract images if multimodal mode
        if extract_images and self.multimodal_embeddings:
            images = self._extract_images_from_pdf(file_path)
            for img_idx, (image, page_num) in enumerate(images):
                # Create a document for each image
                documents.append(Document(
                    page_content=f"[Image {img_idx + 1} from page {page_num}]",
                    metadata={
                        "page": page_num,
                        "has_image": True,
                        "image_index": img_idx,
                        "image_path": image,  # Temporary path to image
                        "content_type": "image"
                    }
                ))
        
        return documents
    
    def _parse_docx(self, file_path: str, extract_images: bool = False) -> List[Document]:
        """Parse DOCX file. Optionally extract images if extract_images=True."""
        doc = docx.Document(file_path)
        documents = []
        current_section = None
        current_subsection = None
        
        # Extract structured content with section/subsection detection
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect section/subsection headers (common patterns)
            section_match = self._detect_section_header(text)
            if section_match:
                if section_match.get("type") == "section":
                    current_section = section_match.get("number")
                    current_subsection = None
                elif section_match.get("type") == "subsection":
                    current_subsection = section_match.get("number")
                
                # Create document for section header
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "section_number": current_section,
                        "subsection_number": current_subsection,
                        "is_header": True
                    }
                ))
            else:
                # Regular paragraph
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "section_number": current_section,
                        "subsection_number": current_subsection
                    }
                ))
        
        # If no sections detected, create single document
        if not documents:
            full_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            documents = [Document(page_content=full_text, metadata={"source": file_path})]
        
        # Extract images if multimodal mode
        if extract_images and self.multimodal_embeddings:
            images = self._extract_images_from_docx(file_path)
            for img_idx, image_path in enumerate(images):
                documents.append(Document(
                    page_content=f"[Image {img_idx + 1} from document]",
                    metadata={
                        "has_image": True,
                        "image_index": img_idx,
                        "image_path": image_path,
                        "content_type": "image"
                    }
                ))
        
        return documents
    
    def _parse_image_file(self, file_path: str, file_type: str) -> List[Document]:
        """Parse standalone image file."""
        return [Document(
            page_content=f"[Image file: {Path(file_path).name}]",
            metadata={
                "has_image": True,
                "image_path": file_path,
                "content_type": "image",
                "file_type": file_type
            }
        )]
    
    def _extract_images_from_pdf(self, file_path: str) -> List[tuple]:
        """Extract images from PDF file. Returns list of (image_path, page_num) tuples."""
        images = []
        try:
            # Try using pdf2image (requires poppler)
            try:
                from pdf2image import convert_from_path
                pages = convert_from_path(file_path, dpi=200)
                temp_dir = Path("./temp/images")
                temp_dir.mkdir(parents=True, exist_ok=True)
                
                for page_num, page_image in enumerate(pages, 1):
                    img_path = temp_dir / f"{Path(file_path).stem}_page_{page_num}.png"
                    page_image.save(img_path, "PNG")
                    images.append((str(img_path), page_num))
            except ImportError:
                # Fallback: try PyMuPDF (fitz)
                try:
                    import fitz  # PyMuPDF
                    pdf_doc = fitz.open(file_path)
                    temp_dir = Path("./temp/images")
                    temp_dir.mkdir(parents=True, exist_ok=True)
                    
                    for page_num in range(len(pdf_doc)):
                        page = pdf_doc[page_num]
                        image_list = page.get_images()
                        
                        for img_idx, img in enumerate(image_list):
                            xref = img[0]
                            base_image = pdf_doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            img_path = temp_dir / f"{Path(file_path).stem}_page_{page_num + 1}_img_{img_idx}.png"
                            
                            with open(img_path, "wb") as f:
                                f.write(image_bytes)
                            images.append((str(img_path), page_num + 1))
                    
                    pdf_doc.close()
                except ImportError:
                    logger.warning("pdf2image and PyMuPDF not available. Install with: pip install pdf2image PyMuPDF")
        except Exception as e:
            logger.warning(f"Failed to extract images from PDF: {e}")
        
        return images
    
    def _extract_images_from_docx(self, file_path: str) -> List[str]:
        """Extract images from DOCX file. Returns list of image file paths."""
        images = []
        try:
            import zipfile
            
            # DOCX is a ZIP file
            docx_zip = zipfile.ZipFile(file_path)
            temp_dir = Path("./temp/images")
            temp_dir.mkdir(parents=True, exist_ok=True)
            
            # Extract images from word/media/ folder
            for file_info in docx_zip.filelist:
                if file_info.filename.startswith('word/media/'):
                    image_data = docx_zip.read(file_info.filename)
                    img_name = Path(file_info.filename).name
                    img_path = temp_dir / f"{Path(file_path).stem}_{img_name}"
                    
                    with open(img_path, "wb") as f:
                        f.write(image_data)
                    images.append(str(img_path))
            
            docx_zip.close()
        except Exception as e:
            logger.warning(f"Failed to extract images from DOCX: {e}")
        
        return images
    
    def _extract_sections_from_text(self, document: Document) -> List[Document]:
        """
        Extract section and subsection numbers from document text.
        Updates document metadata with section/subsection information.
        """
        text = document.page_content
        metadata = document.metadata.copy()
        
        # Detect section patterns in text
        section_match = self._detect_section_header(text)
        if section_match:
            if section_match.get("type") == "section":
                metadata["section_number"] = section_match.get("number")
            elif section_match.get("type") == "subsection":
                metadata["subsection_number"] = section_match.get("number")
                # Also set parent section if detectable
                parts = section_match.get("number", "").split(".")
                if len(parts) >= 2:
                    metadata["section_number"] = ".".join(parts[:-1])
        
        document.metadata = metadata
        return [document]
    
    def _detect_section_header(self, text: str) -> Optional[Dict[str, Any]]:
        """
        Detect section/subsection headers in text.
        Returns dict with type and number if detected.
        
        Patterns:
        - "Section 2.1" or "Section 2" or "Sec. 2.1"
        - "2.1 Introduction" or "2.1.3 Details"
        - "Chapter 3" or "Chapter 3.2"
        """
        import re
        
        # Pattern 1: "Section 2.1" or "Section 2" or "Sec. 2.1"
        pattern1 = r'(?:Section|Sec\.?|§)\s+(\d+(?:\.\d+)*)'
        match = re.search(pattern1, text, re.IGNORECASE)
        if match:
            number = match.group(1)
            dot_count = number.count('.')
            return {
                "type": "subsection" if dot_count >= 2 else ("section" if dot_count >= 1 else "section"),
                "number": number
            }
        
        # Pattern 2: "2.1 Introduction" or "2.1.3 Details" (number at start)
        pattern2 = r'^(\d+(?:\.\d+)+)\s+[A-Z]'
        match = re.search(pattern2, text)
        if match:
            number = match.group(1)
            dot_count = number.count('.')
            return {
                "type": "subsection" if dot_count >= 2 else "section",
                "number": number
            }
        
        # Pattern 3: "Chapter 3" or "Chapter 3.2"
        pattern3 = r'Chapter\s+(\d+(?:\.\d+)*)'
        match = re.search(pattern3, text, re.IGNORECASE)
        if match:
            number = match.group(1)
            return {
                "type": "section",
                "number": number
            }
        
        return None
    
    def _parse_subtitle(self, file_path: str, file_type: str) -> List[Document]:
        """Parse VTT or SRT subtitle files with timestamps."""
        documents = []
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        current_segment = []
        start_time = None
        end_time = None
        
        for line in lines:
            line = line.strip()
            
            # Time format: 00:00:00.000 --> 00:00:00.000
            if '-->' in line:
                times = line.split('-->')
                if len(times) == 2:
                    start_time = times[0].strip()
                    end_time = times[1].strip()
            elif line and not line.isdigit() and not line.startswith('WEBVTT'):
                current_segment.append(line)
            elif not line and current_segment:
                # End of segment
                text = ' '.join(current_segment)
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "start_time": start_time,
                        "end_time": end_time,
                        "source": file_path
                    }
                ))
                current_segment = []
                start_time = None
                end_time = None
        
        # Handle last segment
        if current_segment:
            text = ' '.join(current_segment)
            documents.append(Document(
                page_content=text,
                metadata={
                    "start_time": start_time,
                    "end_time": end_time,
                    "source": file_path
                }
            ))
        
        return documents
    
    def _parse_json_transcript(self, file_path: str) -> List[Document]:
        """Parse structured JSON transcript with speaker turns."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        documents = []
        
        # Handle different JSON structures
        if isinstance(data, list):
            # Array of turns
            for turn in data:
                speaker = turn.get('speaker', 'unknown')
                text = turn.get('text', '')
                start_time = turn.get('start_time')
                end_time = turn.get('end_time')
                
                documents.append(Document(
                    page_content=f"{speaker}: {text}",
                    metadata={
                        "speaker": speaker,
                        "start_time": start_time,
                        "end_time": end_time,
                        "turn_index": turn.get('index', 0)
                    }
                ))
        elif isinstance(data, dict):
            # Object with turns array
            turns = data.get('turns', [])
            for turn in turns:
                speaker = turn.get('speaker', 'unknown')
                text = turn.get('text', '')
                documents.append(Document(
                    page_content=f"{speaker}: {text}",
                    metadata={
                        "speaker": speaker,
                        "start_time": turn.get('start_time'),
                        "end_time": turn.get('end_time')
                    }
                ))
        
        return documents
    
    def _chunk_documents(
        self,
        documents: List[Document],
        file_type: str,
        source_type: str
    ) -> List[Document]:
        """
        Context-aware chunking based on document type.
        
        Strategies:
        - Transcripts (JSON, VTT): Speaker/turn-based
        - Documents: Paragraph-based with semantic boundaries
        - General: Recursive character splitting with overlap
        """
        chunks = []
        
        if file_type == "json" or source_type == "transcript":
            # Speaker/turn-based chunking for transcripts
            for doc in documents:
                # Each turn is already a chunk, but we can combine related turns
                chunks.append(doc)
        elif file_type in ["vtt", "srt"]:
            # Subtitle segments are already chunked
            chunks = documents
        else:
            # Paragraph-based chunking for documents with better definition preservation
            # Use larger chunks and smarter separators to keep definitions together
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=800,  # Increased from 512 to better preserve definitions
                chunk_overlap=100,  # Increased overlap to maintain context
                length_function=len,
                separators=[
                    "\n\n\n",  # Multiple blank lines (section breaks)
                    "\n\n",    # Paragraph breaks
                    ".\n",     # Sentence followed by newline
                    ". ",      # Sentence end
                    "\n",      # Single newline
                    " ",       # Space
                    ""         # Character level
                ]
            )
            
            for doc in documents:
                doc_chunks = text_splitter.split_documents([doc])
                chunks.extend(doc_chunks)
        
        return chunks
    
    def _store_chunks(self, chunks: List[Document], progress_callback: Optional[Callable[[str, float, str], None]] = None) -> int:
        """Generate embeddings and store chunks in vector database."""
        if not chunks:
            logger.warning("No chunks to store")
            return 0
        
        stored = 0
        
        try:
            logger.info(f"Attempting to store chunks using {self.vector_db_type}")
            if self.vector_db_type == "Qdrant" and self.qdrant_client:
                logger.debug("Using Qdrant for storage")
                stored = self._store_in_qdrant(chunks, progress_callback=progress_callback)
            elif self.vector_db_type == "Pinecone" and self.pinecone_index:
                logger.debug("Using Pinecone for storage")
                stored = self._store_in_pinecone(chunks, progress_callback=progress_callback)
            elif hasattr(self, 'chroma_collection') and self.chroma_collection:
                logger.debug("Using ChromaDB for storage")
                stored = self._store_in_chroma(chunks, progress_callback=progress_callback)
            else:
                error_msg = f"No vector database available for storage. DB type: {self.vector_db_type}, Qdrant: {self.qdrant_client is not None}, Chroma: {hasattr(self, 'chroma_collection')}"
                logger.error(error_msg)
                st.error(error_msg)
        except Exception as e:
            error_traceback = traceback.format_exc()
            logger.error(f"Error storing chunks: {str(e)}")
            logger.error(f"Traceback:\n{error_traceback}")
            st.error(f"Error storing chunks: {e}")
            st.exception(e)
            
            # Log to file
            error_log_msg = f"""STORAGE ERROR
Vector DB Type: {self.vector_db_type}
Chunks to store: {len(chunks)}
Error: {str(e)}
Traceback:
{error_traceback}"""
            write_to_log(error_log_msg, "STORAGE ERROR")
        
        return stored
    
    def _get_text_embedding(self, text: str) -> List[float]:
        """Get text embedding using available embedding method."""
        if hasattr(self.embeddings, 'embed_text'):
            # Multimodal embeddings
            return self.embeddings.embed_text(text)
        elif hasattr(self.embeddings, 'embed_query'):
            # LangChain embeddings
            return self.embeddings.embed_query(text)
        else:
            raise ValueError("Embeddings object does not have embed_query or embed_text method")
    
    def _store_in_qdrant(self, chunks: List[Document], progress_callback: Optional[callable] = None) -> int:
        """Store chunks in Qdrant."""
        if self.qdrant_client is None:
            logger.error("Qdrant client is None - cannot store chunks")
            raise AttributeError("Qdrant client not initialized")
        
        collection_name = f"project_{self.project_id}"
        points = []
        total_chunks = len(chunks)
        
        logger.debug(f"Generating embeddings for {total_chunks} chunks...")
        
        # Generate embeddings with progress updates
        for i, chunk in enumerate(chunks):
            try:
                # Update progress
                if progress_callback:
                    progress = 0.70 + (0.25 * (i + 1) / total_chunks)
                    progress_callback(
                        "Generating embeddings",
                        progress,
                        f"Chunk {i + 1}/{total_chunks} ({chunk.metadata.get('file_name', 'unknown')})"
                    )
                
                # Generate embedding - handle both multimodal and text-only
                # Check if this chunk contains an image
                has_image = chunk.metadata.get('has_image', False)
                image_path = chunk.metadata.get('image_path')
                
                if has_image and image_path and hasattr(self.embeddings, 'embed_image'):
                    # Embed image if available and multimodal model supports it
                    try:
                        from PIL import Image
                        if os.path.exists(image_path):
                            img = Image.open(image_path).convert("RGB")
                            embedding = self.embeddings.embed_image(img)
                        else:
                            # Fallback to text embedding if image not found
                            embedding = self._get_text_embedding(chunk.page_content)
                    except Exception as e:
                        logger.warning(f"Failed to embed image {image_path}: {e}, falling back to text")
                        embedding = self._get_text_embedding(chunk.page_content)
                else:
                    # Text embedding
                    embedding = self._get_text_embedding(chunk.page_content)
                
                # Create point
                if PointStruct is not None:
                    point = PointStruct(
                        id=hash(chunk.page_content + str(chunk.metadata)) % (2**63),
                        vector=embedding,
                        payload={
                            "text": chunk.page_content,
                            **chunk.metadata
                        }
                    )
                    points.append(point)
                else:
                    points.append({
                        "id": hash(chunk.page_content + str(chunk.metadata)) % (2**63),
                        "vector": embedding,
                        "payload": {
                            "text": chunk.page_content,
                            **chunk.metadata
                        }
                    })
            except Exception as e:
                logger.error(f"Error generating embedding for chunk {i}: {e}")
                raise
        
        if progress_callback:
            progress_callback("Storing in Qdrant", 0.95, f"Uploading {len(points)} vectors...")
        
        logger.debug(f"Upserting {len(points)} points to Qdrant collection: {collection_name}")
        # Batch upsert
        self.qdrant_client.upsert(
            collection_name=collection_name,
            points=points
        )
        logger.info(f"Successfully stored {len(points)} chunks in Qdrant")
        return len(points)
    
    def _store_in_pinecone(self, chunks: List[Document], progress_callback: Optional[callable] = None) -> int:
        """Store chunks in Pinecone."""
        vectors_to_upsert = []
        total_chunks = len(chunks)
        
        for i, chunk in enumerate(chunks):
            # Update progress
            if progress_callback:
                progress = 0.70 + (0.25 * (i + 1) / total_chunks)
                progress_callback(
                    "Generating embeddings",
                    progress,
                    f"Chunk {i + 1}/{total_chunks}"
                )
            
            # Generate embedding - handle both multimodal and text-only
            if hasattr(self.embeddings, 'embed_text'):
                # Multimodal embeddings
                embedding = self.embeddings.embed_text(chunk.page_content)
            elif hasattr(self.embeddings, 'embed_query'):
                # LangChain embeddings
                embedding = self.embeddings.embed_query(chunk.page_content)
            else:
                raise ValueError("Embeddings object does not have embed_query or embed_text method")
            
            # Create vector for Pinecone
            vector_id = f"{self.project_id}_{hash(chunk.page_content + str(chunk.metadata)) % (2**63)}"
            
            vectors_to_upsert.append({
                "id": str(vector_id),
                "values": embedding,
                "metadata": {
                    "text": chunk.page_content,
                    **chunk.metadata
                }
            })
        
        if progress_callback:
            progress_callback("Storing in Pinecone", 0.95, f"Uploading {len(vectors_to_upsert)} vectors...")
        
        # Batch upsert (Pinecone handles batching)
        self.pinecone_index.upsert(vectors=vectors_to_upsert)
        return len(vectors_to_upsert)
    
    def _store_in_chroma(self, chunks: List[Document], progress_callback: Optional[callable] = None) -> int:
        """Store chunks in ChromaDB."""
        texts = [chunk.page_content for chunk in chunks]
        metadatas = [chunk.metadata for chunk in chunks]
        ids = [f"{self.project_id}_{i}" for i in range(len(chunks))]
        
        if progress_callback:
            progress_callback("Generating embeddings", 0.70, f"Creating embeddings for {len(chunks)} chunks...")
        
        embeddings = self.embeddings.embed_documents(texts)
        
        if progress_callback:
            progress_callback("Storing in ChromaDB", 0.95, f"Uploading {len(chunks)} vectors...")
        
        self.chroma_collection.add(
            embeddings=embeddings,
            documents=texts,
            metadatas=metadatas,
            ids=ids
        )
        return len(chunks)

