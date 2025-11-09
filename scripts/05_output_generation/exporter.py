"""
Export Module
Handles export of research reports to DOCX, PDF, HTML formats.
"""

import os
import logging
from typing import Dict, Optional
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

# Try to import export libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Install with: pip install python-docx")

try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError, Exception) as e:
    # Catch ImportError, OSError (DLL loading issues on Windows), and any other exceptions
    WEASYPRINT_AVAILABLE = False
    logger.warning(f"WeasyPrint not available: {e}. PDF export may be limited.")
    try:
        import pdfkit
        PDFKIT_AVAILABLE = True
    except (ImportError, OSError, Exception):
        PDFKIT_AVAILABLE = False
        logger.warning("PDF export not available. Install WeasyPrint or pdfkit: pip install weasyprint OR pip install pdfkit")

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    logger.warning("markdown not available. Install with: pip install markdown")

try:
    from bs4 import BeautifulSoup
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False
    logger.warning("BeautifulSoup4 not available. Install with: pip install beautifulsoup4")


class ReportExporter:
    """Exports research reports to various formats."""
    
    def __init__(self):
        """Initialize the exporter."""
        self.output_dir = Path("outputs")
        self.output_dir.mkdir(exist_ok=True)
    
    def export_docx(
        self,
        report: Dict,
        output_path: Optional[str] = None
    ) -> str:
        """
        Export report to DOCX format.
        
        Args:
            report: Report dictionary with markdown content
            output_path: Optional output file path
            
        Returns:
            Path to generated DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            query_slug = report.get("metadata", {}).get("query", "research")[:50].replace(" ", "_").replace("?", "")
            output_path = str(self.output_dir / f"research_report_{query_slug}_{timestamp}.docx")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Research Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Research Question
        query = report.get("metadata", {}).get("query", "")
        if query:
            doc.add_paragraph(f"Research Question: {query}")
            doc.add_paragraph(f"Generated: {report.get('metadata', {}).get('timestamp', '')}")
            doc.add_paragraph("")
        
        # Sections
        sections = report.get("sections", {})
        section_order = ["Abstract", "Introduction", "Research Findings", "Conclusion"]
        
        for section_name in section_order:
            if section_name in sections:
                doc.add_heading(section_name, level=1)
                content = sections[section_name]
                # Split into paragraphs
                for para in content.split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
                doc.add_paragraph("")
        
        # References
        doc.add_heading("References", level=1)
        references = report.get("references", [])
        for i, ref in enumerate(references, 1):
            url = ref.get("url", "")
            title = ref.get("title", "")
            if url:
                para = doc.add_paragraph(f"[{i}] ", style='List Number')
                para.add_run(url).bold = False
            elif title:
                doc.add_paragraph(f"[{i}] {title}", style='List Number')
        
        doc.save(output_path)
        logger.info(f"Exported DOCX to {output_path}")
        return output_path
    
    def export_pdf(
        self,
        report: Dict,
        output_path: Optional[str] = None
    ) -> str:
        """
        Export report to PDF format.
        
        Args:
            report: Report dictionary with markdown content
            output_path: Optional output file path
            
        Returns:
            Path to generated PDF file
        """
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            query_slug = report.get("metadata", {}).get("query", "research")[:50].replace(" ", "_").replace("?", "")
            output_path = str(self.output_dir / f"research_report_{query_slug}_{timestamp}.pdf")
        
        # Convert markdown to HTML first
        markdown_content = report.get("markdown", "")
        
        if WEASYPRINT_AVAILABLE:
            # Use WeasyPrint (recommended)
            if MARKDOWN_AVAILABLE:
                try:
                    # Try with extensions, fallback to basic if extensions not available
                    try:
                        html_content = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
                    except:
                        html_content = markdown.markdown(markdown_content)
                except Exception as e:
                    logger.warning(f"Markdown conversion error, using fallback: {e}")
                    html_content = self._markdown_to_html(markdown_content)
            else:
                # Fallback: convert markdown manually
                html_content = self._markdown_to_html(markdown_content)
            
            # Add basic CSS styling
            styled_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
                    h1 {{ color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }}
                    h2 {{ color: #555; margin-top: 30px; }}
                    p {{ margin-bottom: 15px; }}
                    code {{ background-color: #f4f4f4; padding: 2px 5px; border-radius: 3px; }}
                    pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            HTML(string=styled_html).write_pdf(output_path)
            logger.info(f"Exported PDF to {output_path}")
            return output_path
        
        elif PDFKIT_AVAILABLE:
            # Use pdfkit (requires wkhtmltopdf binary)
            if MARKDOWN_AVAILABLE:
                html_content = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
            else:
                html_content = self._markdown_to_html(markdown_content)
            
            styled_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
                    h1 {{ color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }}
                    h2 {{ color: #555; margin-top: 30px; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            pdfkit.from_string(styled_html, output_path)
            logger.info(f"Exported PDF to {output_path}")
            return output_path
        
        else:
            raise ImportError("No PDF library available. Install WeasyPrint or pdfkit: pip install weasyprint OR pip install pdfkit")
    
    def export_html(
        self,
        report: Dict,
        output_path: Optional[str] = None
    ) -> str:
        """
        Export report to HTML format.
        
        Args:
            report: Report dictionary with markdown content
            output_path: Optional output file path
            
        Returns:
            Path to generated HTML file
        """
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            query_slug = report.get("metadata", {}).get("query", "research")[:50].replace(" ", "_").replace("?", "")
            output_path = str(self.output_dir / f"research_report_{query_slug}_{timestamp}.html")
        
        markdown_content = report.get("markdown", "")
        
        if MARKDOWN_AVAILABLE:
            try:
                # Try with extensions, fallback to basic if extensions not available
                try:
                    html_content = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
                except:
                    html_content = markdown.markdown(markdown_content)
            except Exception as e:
                logger.warning(f"Markdown conversion error, using fallback: {e}")
                html_content = self._markdown_to_html(markdown_content)
        else:
            html_content = self._markdown_to_html(markdown_content)
        
        # Create full HTML document with styling
        full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Research Report</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
            background-color: #fff;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            margin-top: 0;
        }}
        h2 {{
            color: #34495e;
            margin-top: 40px;
            border-bottom: 1px solid #ecf0f1;
            padding-bottom: 5px;
        }}
        p {{
            margin-bottom: 15px;
            text-align: justify;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            border-left: 4px solid #3498db;
        }}
        a {{
            color: #3498db;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
        hr {{
            border: none;
            border-top: 2px solid #ecf0f1;
            margin: 30px 0;
        }}
        @media print {{
            body {{
                max-width: 100%;
                padding: 0;
            }}
        }}
    </style>
</head>
<body>
    {html_content}
</body>
</html>
"""
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_html)
        
        logger.info(f"Exported HTML to {output_path}")
        return output_path
    
    def _markdown_to_html(self, markdown_text: str) -> str:
        """Convert markdown to HTML manually (fallback)."""
        # Simple markdown to HTML conversion
        html = markdown_text
        
        # Headers
        html = html.replace("### ", "<h3>").replace("\n###", "</h3>\n###")
        html = html.replace("## ", "<h2>").replace("\n##", "</h2>\n##")
        html = html.replace("# ", "<h1>").replace("\n#", "</h1>\n#")
        
        # Paragraphs
        paragraphs = html.split("\n\n")
        html = "\n".join(f"<p>{p.strip()}</p>" if p.strip() and not p.strip().startswith("<") else p for p in paragraphs)
        
        # Links
        import re
        html = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', html)
        
        # Bold
        html = html.replace("**", "<strong>").replace("**", "</strong>")
        
        # Italic
        html = html.replace("*", "<em>").replace("*", "</em>")
        
        return html

